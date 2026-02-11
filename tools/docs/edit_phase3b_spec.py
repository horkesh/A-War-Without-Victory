#!/usr/bin/env python3
"""
Deterministic edit of Manual and Rulebook .docx: add Phase 3B formal spec
and Rulebook summary. No timestamps, no random IDs, stable ordering.
Invoked by tools/docs/invoke_edit_docx.ts (Node).
"""

from __future__ import annotations

import os
import re
import sys
from pathlib import Path

# Ensure we can import phase3b_spec_text
SCRIPT_DIR = Path(__file__).resolve().parent
sys.path.insert(0, str(SCRIPT_DIR))

from docx import Document
from phase3b_spec_text import (
    MANUAL_SECTION_TITLE,
    MANUAL_SPEC_BODY,
    RULEBOOK_SUBSECTION_BODY,
    RULEBOOK_SUBSECTION_TITLE,
)


def repo_root() -> Path:
    r = SCRIPT_DIR.parent.parent
    return r


def manual_path() -> Path:
    return repo_root() / "docs" / "A_War_Without_Victory_Systems_And_Mechanics_Manual_v0_2_5.docx"


def rulebook_path() -> Path:
    return repo_root() / "docs" / "A_War_Without_Victory_Rulebook_v0_2_5.docx"


def _is_section_header(text: str) -> bool:
    return bool(re.match(r"^\d+(\.\d+)?\.\s+[A-Z]", text.strip()))


def _manual_blocks() -> list[tuple[str, str]]:
    """(text, style). Style: Heading 1, Heading 2, Normal."""
    blocks: list[tuple[str, str]] = []
    raw = MANUAL_SPEC_BODY.strip()
    for chunk in re.split(r"\n\n+", raw):
        chunk = chunk.strip()
        if not chunk:
            continue
        if _is_section_header(chunk):
            style = "Heading 2"
        else:
            style = "Normal"
        blocks.append((chunk, style))
    return blocks


def _manual_has_phase3b(doc: Document) -> bool:
    combined = " ".join(p.text for p in doc.paragraphs)
    return MANUAL_SECTION_TITLE in combined


def edit_manual() -> bool:
    path = manual_path()
    doc = Document(str(path))
    if _manual_has_phase3b(doc):
        return False
    # Find Phase 3A section title to insert Phase 3B immediately after it
    target = None
    for i, p in enumerate(doc.paragraphs):
        if "Phase 3A — Pressure Eligibility and Diffusion (Design Freeze)" in p.text:
            # Find the next Heading 1 or section boundary
            # Look for the next major section (Heading 1) after Phase 3A
            for j in range(i + 1, len(doc.paragraphs)):
                next_p = doc.paragraphs[j]
                if next_p.style and "Heading 1" in getattr(next_p.style, "name", ""):
                    target = next_p
                    break
            # If no Heading 1 found, use "8. Command and control degradation" as fallback
            if target is None:
                for p2 in doc.paragraphs:
                    if "8. Command and control degradation" in p2.text:
                        target = p2
                        break
            break
    if target is None:
        raise SystemExit("Manual: could not find Phase 3A section or '8. Command and control degradation'")

    # Insert Phase 3B section before target paragraph.
    # insert_paragraph_before inserts before target; last-added appears
    # closest to target. So we insert in reverse: blocks first (reversed),
    # then title, so document order is title -> blocks -> target.
    blocks = _manual_blocks()
    for text, style in reversed(blocks):
        target.insert_paragraph_before(text, style=style)
    target.insert_paragraph_before(MANUAL_SECTION_TITLE, style="Heading 1")

    doc.save(str(path))
    return True


def _rulebook_has_phase3b(doc: Document) -> bool:
    combined = " ".join(p.text for p in doc.paragraphs)
    return RULEBOOK_SUBSECTION_TITLE in combined


def edit_rulebook() -> bool:
    path = rulebook_path()
    doc = Document(str(path))
    if _rulebook_has_phase3b(doc):
        return False
    # Find Phase 3A subsection to insert Phase 3B immediately after it
    target = None
    for i, p in enumerate(doc.paragraphs):
        if "Phase 3A — Pressure eligibility and diffusion" in p.text:
            # Find the next Heading 3 or section boundary after Phase 3A subsection
            # Look ahead for the next Heading 3 or major section
            for j in range(i + 1, len(doc.paragraphs)):
                next_p = doc.paragraphs[j]
                if next_p.style and "Heading 3" in getattr(next_p.style, "name", ""):
                    target = next_p
                    break
                # Also check for Heading 2 (major sections)
                if next_p.style and "Heading 2" in getattr(next_p.style, "name", ""):
                    target = next_p
                    break
            # If no heading found, use "Authority and control" as fallback
            if target is None:
                for p2 in doc.paragraphs:
                    if "Authority and control" in p2.text and p2.style and "Heading" in getattr(p2.style, "name", ""):
                        target = p2
                        break
            break
    if target is None:
        raise SystemExit("Rulebook: could not find Phase 3A subsection or 'Authority and control' heading")

    # Insert Phase 3B subsection (H3 + body) before target.
    # insert_paragraph_before: last-added closest to target. Insert body then
    # heading so order is heading -> body -> target.
    target.insert_paragraph_before(RULEBOOK_SUBSECTION_BODY, style="Normal")
    target.insert_paragraph_before(RULEBOOK_SUBSECTION_TITLE, style="Heading 3")

    doc.save(str(path))
    return True


def main() -> None:
    os.chdir(repo_root())
    m = edit_manual()
    r = edit_rulebook()
    if not m and not r:
        pass


if __name__ == "__main__":
    main()
