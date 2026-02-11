#!/usr/bin/env python3
"""
Deterministic validation of Phase 3C doc integration.
- Manual contains exact section title once.
- Rulebook contains subsection title once.
- Rulebook references Manual section title exactly.
- Phase ordering verified (3A → 3B → 3C).
Exit 0 if all pass; non-zero otherwise. No timestamps, no randomness.
"""

from __future__ import annotations

import sys
from pathlib import Path

from docx import Document

SCRIPT_DIR = Path(__file__).resolve().parent
REPO = SCRIPT_DIR.parent.parent
MANUAL_TITLE = "Phase 3C — Exhaustion → Collapse Gating (Design Freeze)"
RULEBOOK_SUBSECTION = "Phase 3C — Exhaustion and collapse eligibility"
PHASE3A_TITLE = "Phase 3A — Pressure Eligibility and Diffusion (Design Freeze)"
PHASE3B_TITLE = "Phase 3B — Pressure → Exhaustion Coupling (Design Freeze)"


def main() -> int:
    manual_path = REPO / "docs" / "A_War_Without_Victory_Systems_And_Mechanics_Manual_v0_2_5.docx"
    rulebook_path = REPO / "docs" / "A_War_Without_Victory_Rulebook_v0_2_5.docx"

    m = Document(str(manual_path))
    r = Document(str(rulebook_path))
    mtext = " ".join(p.text for p in m.paragraphs)
    rtext = " ".join(p.text for p in r.paragraphs)

    ok = True
    n = mtext.count(MANUAL_TITLE)
    if n != 1:
        print(f"FAIL: Manual must contain section title exactly once; found {n}")
        ok = False
    else:
        print("OK: Manual contains section title exactly once")

    n = rtext.count(RULEBOOK_SUBSECTION)
    if n != 1:
        print(f"FAIL: Rulebook must contain subsection title exactly once; found {n}")
        ok = False
    else:
        print("OK: Rulebook contains subsection title exactly once")

    if MANUAL_TITLE not in rtext or "Systems & Mechanics Manual" not in rtext:
        print("FAIL: Rulebook must reference Manual section title exactly")
        ok = False
    else:
        print("OK: Rulebook references Manual section title")

    # Verify phase ordering in Manual
    phase3a_pos = mtext.find(PHASE3A_TITLE)
    phase3b_pos = mtext.find(PHASE3B_TITLE)
    phase3c_pos = mtext.find(MANUAL_TITLE)
    if phase3a_pos == -1 or phase3b_pos == -1 or phase3c_pos == -1:
        print("FAIL: Could not find all Phase 3 sections in Manual")
        ok = False
    elif not (phase3a_pos < phase3b_pos < phase3c_pos):
        print("FAIL: Phase ordering incorrect in Manual (expected 3A -> 3B -> 3C)")
        ok = False
    else:
        print("OK: Phase ordering correct in Manual (3A -> 3B -> 3C)")

    # Verify phase ordering in Rulebook
    phase3a_rulebook = "Phase 3A — Pressure eligibility and diffusion"
    phase3b_rulebook = "Phase 3B — Pressure and exhaustion"
    phase3c_rulebook = RULEBOOK_SUBSECTION
    r_phase3a_pos = rtext.find(phase3a_rulebook)
    r_phase3b_pos = rtext.find(phase3b_rulebook)
    r_phase3c_pos = rtext.find(phase3c_rulebook)
    if r_phase3a_pos == -1 or r_phase3b_pos == -1 or r_phase3c_pos == -1:
        print("FAIL: Could not find all Phase 3 subsections in Rulebook")
        ok = False
    elif not (r_phase3a_pos < r_phase3b_pos < r_phase3c_pos):
        print("FAIL: Phase ordering incorrect in Rulebook (expected 3A -> 3B -> 3C)")
        ok = False
    else:
        print("OK: Phase ordering correct in Rulebook (3A -> 3B -> 3C)")

    return 0 if ok else 1


if __name__ == "__main__":
    sys.exit(main())
