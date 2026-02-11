"""
Apply minimal AoR scoping reconciliation edits to non-Rulebook docs.

Edits:
- Systems & Mechanics Manual v0.2.5: replace "every settlement must be assigned..." invariant
- Game Bible v0.2.5: replace "one brigade per settlement ... no overlap and no gaps" invariant
- Engine Invariants v0.2.6: update only if it still asserts universal AoR assignment

Determinism:
- stable file ordering
- no timestamps
- no random IDs
"""

from __future__ import annotations

import re
from pathlib import Path
from typing import Iterable, Optional

from docx import Document  # type: ignore
from docx.oxml import OxmlElement  # type: ignore
from docx.text.paragraph import Paragraph  # type: ignore


ROOT = Path(__file__).resolve().parents[2]
DOCS = ROOT / "docs"

RULEBOOK = DOCS / "A_War_Without_Victory_Rulebook_v0_2_6.docx"

SYSTEMS_MANUAL = DOCS / "A_War_Without_Victory_Systems_And_Mechanics_Manual_v0_2_5.docx"
GAME_BIBLE = DOCS / "A_War_Without_Victory_The_Game_Bible_v0_2_5.docx"
ENGINE_INVARIANTS = DOCS / "A_War_Without_Victory_Engine_Invariants_v0_2_6.docx"

SYSTEMS_MANUAL_V023 = DOCS / "A_War_Without_Victory_Systems_And_Mechanics_Manual_v0_2_3_FINAL.docx"
GAME_BIBLE_V023 = DOCS / "A_War_Without_Victory_The_Game_Bible_v0_2_3_FINAL.docx"


def insert_paragraph_after(paragraph: Paragraph, text: str, style: Optional[str] = None) -> Paragraph:
    """Insert a new paragraph after `paragraph` with plain text (no numbering)."""
    new_p = OxmlElement("w:p")
    paragraph._p.addnext(new_p)  # type: ignore[attr-defined]
    new_para = Paragraph(new_p, paragraph._parent)  # type: ignore[arg-type]
    if style:
        try:
            new_para.style = style
        except Exception:
            pass
    new_para.add_run(text)
    return new_para


def extract_rulebook_canon_lines() -> list[str]:
    """
    Extract the AoR/control canon lines from Rulebook v0.2.6.
    We take the AoR section heading and all paragraphs until next heading.
    """
    if not RULEBOOK.exists():
        raise FileNotFoundError(str(RULEBOOK))

    doc = Document(str(RULEBOOK))
    lines: list[str] = []
    in_section = False
    for p in doc.paragraphs:
        txt = (p.text or "").strip()
        if not txt:
            continue
        style_name = ""
        try:
            style_name = (p.style.name or "").strip()
        except Exception:
            style_name = ""
        is_heading = style_name.lower().startswith("heading")
        if is_heading:
            if re.search(r"\bareas of responsibility\b|\bAoR\b", txt, re.IGNORECASE):
                in_section = True
                lines.append(txt)
                continue
            if in_section:
                break
        if in_section:
            lines.append(txt)
    if not lines:
        raise RuntimeError("Could not deterministically find AoR section heading in Rulebook v0.2.6")
    return lines


def canon_scope_note_lines() -> list[str]:
    """
    Minimal subset, verbatim from Rulebook v0.2.6 AoR section, used as a scope note.
    """
    canon = extract_rulebook_canon_lines()

    wanted_starts = [
        "AoRs apply only to front-active settlements",
        "Settlements that are politically controlled",
        "Rear Political Control Zones",
        "Settlement control does not change due to lack of brigade presence.",
        "Control change may occur only when:",
        "Once front-active, the settlement must be assigned to exactly one brigade",
    ]

    out: list[str] = []
    for line in canon:
        if any(line.startswith(ws) for ws in wanted_starts):
            out.append(line)
        # Include the two bullet lines that follow "Control change may occur only when:"
        if out and out[-1].startswith("Control change may occur only when:"):
            continue
        # Include the bullet lines that follow the control-change lead line
        if out and len(out) >= 1 and out[-1].startswith("Control change may occur only when:"):
            pass

    # The above is conservative; add the immediate bullets for control-change if present in canon.
    # We include the two lines following the exact "Control change may occur only when:" match.
    for i, line in enumerate(canon):
        if line == "Control change may occur only when:":
            for j in range(i + 1, min(i + 3, len(canon))):
                out.append(canon[j])
            break

    # If we captured only headings without content, fall back to a fixed slice:
    if len(out) < 4:
        # Keep ordering deterministic and minimal: take the last 12 lines in the AoR section.
        out = canon[-12:]

    # De-dup while preserving order
    seen = set()
    deduped: list[str] = []
    for ln in out:
        if ln in seen:
            continue
        seen.add(ln)
        deduped.append(ln)
    return deduped


def replace_paragraph_invariant_with_scope_note(doc_path: Path, match_predicate, scope_note_lines: list[str]) -> bool:
    doc = Document(str(doc_path))
    replaced = False
    for p in doc.paragraphs:
        txt = (p.text or "").strip()
        if not txt:
            continue
        if match_predicate(txt):
            # Minimal edit: replace the conflicting invariant paragraph with a scope note + verbatim lines.
            p.text = "Scope note (Rulebook v0.2.6):"
            insert_after = p
            for ln in scope_note_lines:
                insert_after = insert_paragraph_after(insert_after, ln)
            replaced = True
            break
    if replaced:
        doc.save(str(doc_path))
    return replaced


def main() -> int:
    scope_lines = canon_scope_note_lines()

    targets = [
        SYSTEMS_MANUAL,
        GAME_BIBLE,
        ENGINE_INVARIANTS,
        SYSTEMS_MANUAL_V023,
        GAME_BIBLE_V023,
    ]

    for path in targets:
        if not path.exists():
            raise FileNotFoundError(str(path))

    # Systems manual: replace universal assignment invariant paragraph.
    systems_changed = replace_paragraph_invariant_with_scope_note(
        SYSTEMS_MANUAL,
        lambda t: "Every settlement must be assigned to exactly one brigade at all times." in t
        or ("Assignment may not overlap" in t and "may not be empty" in t),
        scope_lines,
    )

    systems_v023_changed = replace_paragraph_invariant_with_scope_note(
        SYSTEMS_MANUAL_V023,
        lambda t: "Every settlement must be assigned to exactly one brigade at all times." in t
        or ("Assignment may not overlap" in t and "may not be empty" in t),
        scope_lines,
    )

    # Game bible: replace "one brigade per settlement ... no overlap and no gaps" invariant line.
    bible_changed = replace_paragraph_invariant_with_scope_note(
        GAME_BIBLE,
        lambda t: t.lower().startswith("one brigade per settlement:")
        and ("no overlap" in t.lower() and "no gaps" in t.lower()),
        scope_lines,
    )

    bible_v023_changed = replace_paragraph_invariant_with_scope_note(
        GAME_BIBLE_V023,
        lambda t: t.lower().startswith("one brigade per settlement:")
        and ("no overlap" in t.lower() and "no gaps" in t.lower()),
        scope_lines,
    )

    # Engine invariants: only if it still asserts universal assignment.
    invariants_changed = replace_paragraph_invariant_with_scope_note(
        ENGINE_INVARIANTS,
        lambda t: ("every settlement must be assigned" in t.lower())
        or ("one brigade per settlement" in t.lower() and "at all times" in t.lower()),
        scope_lines,
    )

    print("APPLY: AoR doc reconciliation")
    print(f"- {SYSTEMS_MANUAL.name}: {'UPDATED' if systems_changed else 'no change (already compatible / no match)'}")
    print(f"- {SYSTEMS_MANUAL_V023.name}: {'UPDATED' if systems_v023_changed else 'no change (already compatible / no match)'}")
    print(f"- {GAME_BIBLE.name}: {'UPDATED' if bible_changed else 'no change (already compatible / no match)'}")
    print(f"- {GAME_BIBLE_V023.name}: {'UPDATED' if bible_v023_changed else 'no change (already compatible / no match)'}")
    print(f"- {ENGINE_INVARIANTS.name}: {'UPDATED' if invariants_changed else 'no change (already compatible / no match)'}")
    return 0


if __name__ == "__main__":
    raise SystemExit(main())

