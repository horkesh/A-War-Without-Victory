#!/usr/bin/env python3
"""
Deterministic validation of Phase 3B doc integration.
- Manual contains exact section title once.
- Rulebook contains subsection title once.
- Rulebook references Manual section title exactly.
Exit 0 if all pass; non-zero otherwise. No timestamps, no randomness.
"""

from __future__ import annotations

import sys
from pathlib import Path

from docx import Document

SCRIPT_DIR = Path(__file__).resolve().parent
REPO = SCRIPT_DIR.parent.parent
MANUAL_TITLE = "Phase 3B — Pressure → Exhaustion Coupling (Design Freeze)"
RULEBOOK_SUBSECTION = "Phase 3B — Pressure and exhaustion"


def main() -> int:
    manual_path = REPO / "docs" / "A_War_Without_Victory_Systems_And_Mechanics_Manual_v0_2_5.docx"
    rulebook_path = REPO / "docs" / "A_War_Without_Victory_Rulebook_v0_2_5.docx"

    m = Document(str(manual_path))
    r = Document(str(rulebook_path))
    mtext = " ".join(p.text for p in m.paragraphs)
    rtext = " ".join(p.text for p in r.paragraphs)

    ok = True
    n = mtext.count(MANUAL_TITLE)
    if n != 1:
        print(f"FAIL: Manual must contain section title exactly once; found {n}")
        ok = False
    else:
        print("OK: Manual contains section title exactly once")

    n = rtext.count(RULEBOOK_SUBSECTION)
    if n != 1:
        print(f"FAIL: Rulebook must contain subsection title exactly once; found {n}")
        ok = False
    else:
        print("OK: Rulebook contains subsection title exactly once")

    if MANUAL_TITLE not in rtext or "Systems & Mechanics Manual" not in rtext:
        print("FAIL: Rulebook must reference Manual section title exactly")
        ok = False
    else:
        print("OK: Rulebook references Manual section title")

    return 0 if ok else 1


if __name__ == "__main__":
    sys.exit(main())
