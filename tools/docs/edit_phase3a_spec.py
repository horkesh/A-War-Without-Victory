#!/usr/bin/env python3
"""
Deterministic edit of Manual and Rulebook .docx: add Phase 3A formal spec
and Rulebook summary. No timestamps, no random IDs, stable ordering.
Invoked by tools/docs/invoke_edit_docx.ts (Node).
"""

from __future__ import annotations

import os
import re
import sys
from pathlib import Path

# Ensure we can import phase3a_spec_text
SCRIPT_DIR = Path(__file__).resolve().parent
sys.path.insert(0, str(SCRIPT_DIR))

from docx import Document
from phase3a_spec_text import (
    MANUAL_SECTION_TITLE,
    MANUAL_SPEC_BODY,
    RULEBOOK_SUBSECTION_BODY,
    RULEBOOK_SUBSECTION_TITLE,
)


def repo_root() -> Path:
    r = SCRIPT_DIR.parent.parent
    return r


def manual_path() -> Path:
    return repo_root() / "docs" / "A_War_Without_Victory_Systems_And_Mechanics_Manual_v0_2_5.docx"


def rulebook_path() -> Path:
    return repo_root() / "docs" / "A_War_Without_Victory_Rulebook_v0_2_5.docx"


def _is_section_header(text: str) -> bool:
    return bool(re.match(r"^\d+(\.\d+)?\.\s+[A-Z]", text.strip()))


def _manual_blocks() -> list[tuple[str, str]]:
    """(text, style). Style: Heading 1, Heading 2, Normal."""
    blocks: list[tuple[str, str]] = []
    raw = MANUAL_SPEC_BODY.strip()
    for chunk in re.split(r"\n\n+", raw):
        chunk = chunk.strip()
        if not chunk:
            continue
        if _is_section_header(chunk):
            style = "Heading 2"
        else:
            style = "Normal"
        blocks.append((chunk, style))
    return blocks


def _manual_has_phase3a(doc: Document) -> bool:
    combined = " ".join(p.text for p in doc.paragraphs)
    return MANUAL_SECTION_TITLE in combined


def edit_manual() -> bool:
    path = manual_path()
    doc = Document(str(path))
    if _manual_has_phase3a(doc):
        return False
    target = None
    for p in doc.paragraphs:
        if "8. Command and control degradation" in p.text:
            target = p
            break
    if target is None:
        raise SystemExit("Manual: could not find '8. Command and control degradation'")

    # Insert Phase 3A section before that paragraph.
    # insert_paragraph_before inserts before target; last-added appears
    # closest to target. So we insert in reverse: blocks first (reversed),
    # then title, so document order is title -> blocks -> "8. Command...".
    blocks = _manual_blocks()
    for text, style in reversed(blocks):
        target.insert_paragraph_before(text, style=style)
    target.insert_paragraph_before(MANUAL_SECTION_TITLE, style="Heading 1")

    doc.save(str(path))
    return True


def _rulebook_has_phase3a(doc: Document) -> bool:
    combined = " ".join(p.text for p in doc.paragraphs)
    return RULEBOOK_SUBSECTION_TITLE in combined


def edit_rulebook() -> bool:
    path = rulebook_path()
    doc = Document(str(path))
    if _rulebook_has_phase3a(doc):
        return False
    target = None
    for p in doc.paragraphs:
        if "Authority and control" in p.text and p.style and "Heading" in getattr(p.style, "name", ""):
            target = p
            break
    if target is None:
        raise SystemExit("Rulebook: could not find 'Authority and control' heading")

    # Insert Phase 3A subsection (H3 + body) before "Authority and control".
    # insert_paragraph_before: last-added closest to target. Insert body then
    # heading so order is heading -> body -> "Authority and control".
    target.insert_paragraph_before(RULEBOOK_SUBSECTION_BODY, style="Normal")
    target.insert_paragraph_before(RULEBOOK_SUBSECTION_TITLE, style="Heading 3")

    doc.save(str(path))
    return True


def main() -> None:
    os.chdir(repo_root())
    m = edit_manual()
    r = edit_rulebook()
    if not m and not r:
        pass


if __name__ == "__main__":
    main()
