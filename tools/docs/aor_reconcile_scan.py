"""
Scan AWWV docs for AoR scoping contradictions vs Rulebook v0.2.6.

Determinism:
- stable sorted file ordering
- no timestamps emitted

Output:
- "CANON:" block extracted from Rulebook v0.2.6 around AoR section (best-effort)
- "CONFLICTS:" list (filename, last heading, paragraph text) for other docx/md files

This script is read-only: it does not modify any documents.
"""

from __future__ import annotations

import os
import re
from pathlib import Path
from typing import Iterable, Optional, Tuple

from docx import Document  # type: ignore


ROOT = Path(__file__).resolve().parents[2]
DOCS_DIR = ROOT / "docs"


RULEBOOK_V026 = DOCS_DIR / "A_War_Without_Victory_Rulebook_v0_2_6.docx"


CONFLICT_PATTERNS = [
    re.compile(r"\beach settlement\b.*\bassigned\b", re.IGNORECASE),
    re.compile(r"\bevery settlement\b.*\bmust be assigned\b", re.IGNORECASE),
    re.compile(r"\bno overlap\b.*\bno gaps\b", re.IGNORECASE),
    re.compile(r"\bunassigned settlement\b.*\binvalid\b", re.IGNORECASE),
    re.compile(r"\bassigned to exactly one\b.*\bAoR\b", re.IGNORECASE),
    re.compile(r"\bAoR\b.*\binvariants remain\b", re.IGNORECASE),
]


def _iter_docx_paragraphs(doc_path: Path) -> Iterable[Tuple[str, str]]:
    """
    Yield (heading, paragraph_text) in order.
    Heading is "best effort": last-seen paragraph with Heading style.
    """
    doc = Document(str(doc_path))
    heading = ""
    for p in doc.paragraphs:
        text = (p.text or "").strip()
        if not text:
            continue
        style_name = ""
        try:
            style_name = (p.style.name or "").strip()
        except Exception:
            style_name = ""
        if style_name.lower().startswith("heading"):
            heading = text
            continue
        yield (heading, text)


def _extract_rulebook_aor_snippet() -> list[str]:
    """
    Extract a best-effort snippet around the AoR section from Rulebook v0.2.6.
    We look for a heading containing "Areas of Responsibility" or "AoR", then
    include subsequent non-empty paragraphs until the next Heading.
    """
    if not RULEBOOK_V026.exists():
        return [f"(missing) {RULEBOOK_V026}"]

    doc = Document(str(RULEBOOK_V026))
    lines: list[str] = []
    in_section = False
    for p in doc.paragraphs:
        text = (p.text or "").strip()
        if not text:
            continue
        style_name = ""
        try:
            style_name = (p.style.name or "").strip()
        except Exception:
            style_name = ""

        is_heading = style_name.lower().startswith("heading")
        if is_heading:
            # Start condition
            if re.search(r"\bareas of responsibility\b|\bAoR\b", text, re.IGNORECASE):
                in_section = True
                lines.append(text)
                continue
            # End condition
            if in_section:
                break

        if in_section:
            lines.append(text)

    if not lines:
        # Fallback: grab first N paragraphs containing AoR-ish terms
        hits: list[str] = []
        for _heading, para in _iter_docx_paragraphs(RULEBOOK_V026):
            if re.search(r"\bAoR\b|\bAreas of Responsibility\b|\bfront-active\b|\brear\b", para, re.IGNORECASE):
                hits.append(para)
            if len(hits) >= 12:
                break
        if hits:
            return ["(fallback snippet; AoR heading not found deterministically)"] + hits
        return ["(no AoR snippet found)"]
    return lines


def _iter_md_conflicts(md_path: Path) -> Iterable[Tuple[str, str]]:
    text = md_path.read_text(encoding="utf-8", errors="replace")
    heading = ""
    for raw in text.splitlines():
        line = raw.rstrip("\n")
        m = re.match(r"^\s*(#{1,6})\s+(.*\S)\s*$", line)
        if m:
            heading = m.group(2).strip()
            continue
        if not line.strip():
            continue
        for pat in CONFLICT_PATTERNS:
            if pat.search(line):
                yield (heading, line.strip())
                break


def main() -> int:
    print("CANON: Rulebook v0.2.6 AoR/control snippet (best-effort)")
    for ln in _extract_rulebook_aor_snippet():
        print(f"- {ln}")
    print()

    docx_paths = sorted([p for p in DOCS_DIR.glob("*.docx") if p.name != RULEBOOK_V026.name], key=lambda p: p.name)
    md_paths = sorted([p for p in DOCS_DIR.glob("*.md")], key=lambda p: p.name)

    conflicts: list[Tuple[str, str, str]] = []

    for p in docx_paths:
        for heading, para in _iter_docx_paragraphs(p):
            if any(pat.search(para) for pat in CONFLICT_PATTERNS):
                conflicts.append((p.name, heading or "(no heading detected)", para))

    for p in md_paths:
        for heading, line in _iter_md_conflicts(p):
            conflicts.append((p.name, heading or "(no heading detected)", line))

    print("CONFLICTS: potential old AoR/global-assignment claims")
    if not conflicts:
        print("(none found by current patterns)")
        return 0

    for fname, heading, text in sorted(conflicts, key=lambda t: (t[0], t[1], t[2])):
        print(f"- {fname} | {heading} | {text}")
    return 0


if __name__ == "__main__":
    raise SystemExit(main())

